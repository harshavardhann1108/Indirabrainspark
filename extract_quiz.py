from docx import Document
import json

# Read the document
doc = Document('Inetractive Quiz.docx')

# Extract all text
all_text = []
for para in doc.paragraphs:
    if para.text.strip():
        all_text.append(para.text)

# Also extract tables if any
tables_data = []
for table in doc.tables:
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        table_data.append(row_data)
    tables_data.append(table_data)

# Save to file
with open('extracted_content.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(all_text))
    f.write('\n\n=== TABLES ===\n\n')
    for i, table in enumerate(tables_data):
        f.write(f'\n--- Table {i+1} ---\n')
        for row in table:
            f.write(' | '.join(row) + '\n')

print("Content extracted successfully!")
print(f"Total paragraphs: {len(all_text)}")
print(f"Total tables: {len(tables_data)}")
